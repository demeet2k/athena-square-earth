# CRYSTAL: Xi108:W2:A4:S28 | face=F | node=398 | depth=2 | phase=Mutable
# METRO: Me,T
# BRIDGES: Xi108:W2:A4:S27→Xi108:W2:A4:S29→Xi108:W1:A4:S28→Xi108:W3:A4:S28→Xi108:W2:A3:S28→Xi108:W2:A5:S28

"""AtlasBook compiler: turn the memory bank into a coherent document.

This is the "next" deepening step for AtlasForge as a math memory bank:

- Memory entries are great for *incremental* work.
- But when you want the *full understanding* to be readable, you want to
  compile those entries into an "atlas" (a book-like artifact):

  * grouped by kind (definitions / lemmas / theorems / operators / ...)
  * sorted by crystal address when available
  * with a table of contents, tag index, and (optionally) graph edges

This module produces:
- Markdown (always available)
- PDF (optional; requires `reportlab`, which is present in the sandbox)
"""

from __future__ import annotations

from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Sequence, Tuple, Union

import os

from atlasforge.memory.entry import MemoryEntry
from atlasforge.memory.store import MemoryStore

def _utc_now_iso() -> str:
    return datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%SZ")

_KIND_ORDER = [
    "seed",
    "operator",
    "identity",
    "definition",
    "lemma",
    "proposition",
    "corollary",
    "theorem",
    "proof",
    "derivation",
    "computation",
    "example",
    "recipe_log",
    "note",
    "other",
]

def _kind_of(entry: MemoryEntry) -> str:
    k = str((entry.extra or {}).get("kind") or "note").strip().lower()
    return k or "note"

def _crystal_index_of(entry: MemoryEntry) -> Optional[int]:
    try:
        v = (entry.extra or {}).get("crystal_index")
        if v is None:
            return None
        return int(v)
    except Exception:
        return None

def _crystal_addr_of(entry: MemoryEntry) -> Optional[str]:
    v = (entry.extra or {}).get("crystal_address") or (entry.extra or {}).get("address")
    if v is None:
        return None
    s = str(v).strip()
    return s or None

def _sanitize_entry_body(entry: MemoryEntry) -> str:
    """Remove redundant first heading if it duplicates the title."""
    content = (entry.content or "").rstrip() + "\n"
    lines = content.splitlines()
    if not lines:
        return content

    # If content starts with '# <title>' then drop it.
    if lines[0].startswith("# "):
        head = lines[0][2:].strip()
        if head and entry.title and head.lower() == entry.title.strip().lower():
            # Drop the header line and one optional blank line.
            lines = lines[1:]
            if lines and not lines[0].strip():
                lines = lines[1:]
            return "\n".join(lines).rstrip() + "\n"
    return content

def _require_all_tags(entry: MemoryEntry, req: Sequence[str]) -> bool:
    if not req:
        return True
    tags = {t.strip().lower() for t in (entry.tags or []) if t.strip()}
    return all(t.strip().lower() in tags for t in req if t.strip())

def _match_query(entry: MemoryEntry, query: str) -> bool:
    q = (query or "").strip().lower()
    if not q:
        return True
    hay = f"{entry.title}\n{entry.content}".lower()
    return q in hay

@dataclass
class AtlasBookConfig:
    """Config for building a book from a MemoryStore."""

    title: str = "AtlasForge Memory Atlas"
    subtitle: str = ""
    created_at: str = field(default_factory=_utc_now_iso)

    # Selection
    query: str = ""
    tags: List[str] = field(default_factory=list)
    kinds: List[str] = field(default_factory=list)  # OR semantics
    session_id: Optional[str] = None

    # Dependency selection (optional)
    root_hashes: List[str] = field(default_factory=list)
    include_dependencies: bool = False
    dependency_relations: List[str] = field(default_factory=lambda: ["depends_on"])
    dependency_direction: str = "out"  # out | in | both
    dependency_max_depth: int = 25

    # Layout
    sort_by: str = "crystal"  # crystal | time | title
    include_crystal_map: bool = True
    include_full_crystal_map: bool = False
    max_entries_per_cell: int = 25
    include_toc: bool = True
    include_tag_index: bool = True
    include_graph_edges: bool = True
    max_entries: Optional[int] = None

class AtlasBookBuilder:
    """Compile MemoryStore entries into Markdown/PDF."""

    def __init__(self, store: MemoryStore):
        self.store = store
        self._sort_by = "crystal"

    # ------------------------------------------------------------------
    # Collection
    # ------------------------------------------------------------------

    def collect(self, config: AtlasBookConfig) -> List[MemoryEntry]:
        """Collect entries matching the config.

        Selection precedence:
        1) ``session_id`` (explicit session export)
        2) ``root_hashes`` (explicit closure export)
        3) query/tags/kinds scan (default)
        """

        entries: List[MemoryEntry] = []

        req_tags = [t.strip().lower() for t in (config.tags or []) if t.strip()]
        allowed_kinds = {k.strip().lower() for k in (config.kinds or []) if k.strip()}

        def accept(e: MemoryEntry) -> bool:
            if req_tags and not _require_all_tags(e, req_tags):
                return False
            if allowed_kinds and _kind_of(e) not in allowed_kinds:
                return False
            if config.query and not _match_query(e, config.query):
                return False
            return True

        # 1) Session selection overrides everything.
        if config.session_id:
            sess = self.store.sessions.get(config.session_id)
            if sess is None:
                return []
            for h in sess.entry_hashes:
                e = self.store.get(h)
                if e is not None and accept(e):
                    entries.append(e)
                    if config.max_entries is not None and len(entries) >= config.max_entries:
                        break
            return entries

        # 2) Root selection: export a dependency-closed subgraph.
        roots = [str(h).strip() for h in (config.root_hashes or []) if str(h).strip()]
        if roots:
            rels = [str(r).strip() for r in (config.dependency_relations or []) if str(r).strip()] or None
            if config.include_dependencies:
                ordered_hashes = self.store.closure(
                    roots,
                    relations=rels,
                    direction=config.dependency_direction,
                    max_depth=int(config.dependency_max_depth or 25),
                    limit=5000,
                )
            else:
                ordered_hashes = roots

            for h in ordered_hashes:
                e = self.store.get(h)
                if e is not None and accept(e):
                    entries.append(e)
                    if config.max_entries is not None and len(entries) >= config.max_entries:
                        break
            return entries

        # 3) Default: scan memory bank and filter.
        for e in self.store.iter_entries():
            if accept(e):
                entries.append(e)
                if config.max_entries is not None and len(entries) >= config.max_entries:
                    break
        return entries

    
    def group(self, entries: Sequence[MemoryEntry]) -> Dict[str, List[MemoryEntry]]:
        """Group entries by kind and sort within each group."""
        groups: Dict[str, List[MemoryEntry]] = {}
        for e in entries:
            groups.setdefault(_kind_of(e), []).append(e)

        # Sort within each kind.
        for _, items in groups.items():
            if not items:
                continue

            def key(e: MemoryEntry):
                if self._sort_by == "title":
                    return (e.title.lower(), e.content_hash())
                if self._sort_by == "time":
                    return (e.created_at, e.content_hash())
                # default: crystal
                idx = _crystal_index_of(e)
                return (idx if idx is not None else 10_000, e.title.lower(), e.content_hash())

            items.sort(key=key)

        return groups
    # ------------------------------------------------------------------
    # Markdown build
    # ------------------------------------------------------------------

    def build_markdown(self, config: AtlasBookConfig) -> str:
        entries = self.collect(config)

        # Determine sort mode.
        sort_by = (config.sort_by or "crystal").strip().lower()
        self._sort_by = sort_by

        groups = self.group(entries)

        # Order kinds.
        ordered_kinds = [k for k in _KIND_ORDER if k in groups]
        # Add any unknown kinds at end.
        for k in sorted(groups.keys()):
            if k not in ordered_kinds:
                ordered_kinds.append(k)

        # Helper: short hash
        def short(h: str) -> str:
            return h[:10] if h else ""

        lines: List[str] = []
        lines.append(f"# {config.title}")
        if config.subtitle:
            lines.append(config.subtitle)
        lines.append("")
        lines.append(f"*Generated:* {config.created_at}")
        if config.session_id:
            lines.append(f"*Session:* {config.session_id}")
        if config.query:
            lines.append(f"*Query:* `{config.query}`")
        if config.tags:
            lines.append(f"*Tags:* {', '.join(config.tags)}")
        if config.kinds:
            lines.append(f"*Kinds:* {', '.join(config.kinds)}")
        if config.root_hashes:
            roots_short = ", ".join([f"`{str(h)[:10]}`" for h in config.root_hashes if str(h)])
            lines.append(f"*Roots:* {roots_short}")
            if config.include_dependencies:
                rels = ", ".join([str(r) for r in (config.dependency_relations or []) if str(r)])
                lines.append(f"*Dependency closure:* relations=[{rels}], direction={config.dependency_direction}, max_depth={config.dependency_max_depth}")
        lines.append("")

        # TOC
        if config.include_toc:
            lines.append("## Table of contents")
            lines.append("")
            for k in ordered_kinds:
                title = k.replace("_", " ").title()
                lines.append(f"- [{title}](#{k})")
                for e in groups.get(k, []):
                    h = e.content_hash()
                    t = e.title or h
                    lines.append(f"  - [{t}](#{h})")
            lines.append("")

        # Crystal navigator / 256-cell index
        if config.include_crystal_map:
            try:
                import atlasforge as af  # local import to avoid heavy import chains
            except Exception:
                af = None  # type: ignore

            cell_map: Dict[int, List[MemoryEntry]] = {}
            unassigned: List[MemoryEntry] = []
            for e in entries:
                idx = _crystal_index_of(e)
                if idx is None:
                    unassigned.append(e)
                else:
                    cell_map.setdefault(int(idx), []).append(e)

            # Stable ordering within cells
            for idx, es in cell_map.items():
                es.sort(key=lambda ee: ((ee.title or "").lower(), ee.content_hash()))

            lines.append("## Crystal map")
            lines.append("")
            lines.append("This section indexes entries by their 4×4×4×4 *crystal address* (Pole·Lens·Layer·Depth).")
            lines.append("")

            # Build a compact 16×16 grid of counts (Pole×Lens vs Layer×Depth).
            if af is not None:
                try:
                    poles = list(af.Pole)
                    lenses = list(af.Lens)
                    layers = list(af.Layer)
                    depths = list(af.Depth)

                    layer_abbr = {
                        "OBJECTS": "Obj",
                        "OPERATORS": "Op",
                        "INVARIANTS": "Inv",
                        "ARTIFACTS": "Art",
                    }

                    col_specs: List[Tuple[int, int, str]] = []
                    for layer in layers:
                        for depth in depths:
                            col_specs.append((layer.value, depth.value, f"{layer_abbr.get(layer.name, layer.name[:3])}{depth.value}"))

                    header = ["Pole/Lens"] + [lab for _, _, lab in col_specs]
                    lines.append("### Crystal grid (counts)")
                    lines.append("")
                    lines.append("| " + " | ".join(header) + " |")
                    lines.append("|" + " --- |" * len(header))

                    for p_i, p in enumerate(poles):
                        for l_i, l in enumerate(lenses):
                            row_label = f"{p.value}{l.value}"
                            row: List[str] = [row_label]
                            for layer_val, depth_val, _lab in col_specs:
                                idx = p_i * 64 + l_i * 16 + int(layer_val) * 4 + int(depth_val)
                                cnt = len(cell_map.get(idx, []))
                                if cnt > 0:
                                    row.append(f"[{cnt}](#cell-{idx})")
                                else:
                                    row.append("0")
                            lines.append("| " + " | ".join(row) + " |")
                    lines.append("")
                except Exception:
                    # Grid is nice-to-have; skip on any failure.
                    pass

            # Detailed cell listing
            lines.append("### Crystal cells")
            lines.append("")
            indices = list(range(256)) if config.include_full_crystal_map else sorted(cell_map.keys())
            for idx in indices:
                es = cell_map.get(int(idx), [])
                if (not es) and (not config.include_full_crystal_map):
                    continue
                addr = str(idx)
                if af is not None:
                    try:
                        addr = str(af.CrystalAddress.from_index(int(idx)))
                    except Exception:
                        addr = str(idx)

                lines.append(f"<a id=\"cell-{idx}\"></a>")
                lines.append(f"#### Cell {idx}: {addr}")
                lines.append("")
                if not es:
                    lines.append("*empty*")
                    lines.append("")
                    continue

                max_cell = int(config.max_entries_per_cell or 25)
                shown = es[:max_cell]
                for e in shown:
                    h = e.content_hash()
                    t = e.title or h
                    lines.append(f"- [{t}](#{h}) (`{h[:10]}`)")
                if len(es) > max_cell:
                    lines.append(f"- … and {len(es) - max_cell} more")
                lines.append("")

            if unassigned:
                lines.append("### Unassigned entries")
                lines.append("")
                for e in unassigned[: min(len(unassigned), 50)]:
                    h = e.content_hash()
                    t = e.title or h
                    lines.append(f"- [{t}](#{h}) (`{h[:10]}`)")
                if len(unassigned) > 50:
                    lines.append(f"- … and {len(unassigned) - 50} more")
                lines.append("")

        # Sections
        tag_index: Dict[str, List[Tuple[str, str]]] = {}

        for k in ordered_kinds:
            section_title = k.replace("_", " ").title()
            lines.append(f"## {section_title}")
            lines.append(f"<a id=\"{k}\"></a>")
            lines.append("")

            for e in groups.get(k, []):
                h = e.content_hash()
                title = e.title or h

                # Tag index build
                for t in (e.tags or []):
                    tt = str(t).strip()
                    if not tt:
                        continue
                    tag_index.setdefault(tt, []).append((title, h))

                lines.append(f"<a id=\"{h}\"></a>")
                lines.append(f"### {title}")

                # Metadata
                meta: List[str] = []
                meta.append(f"- hash: `{h}`")
                meta.append(f"- kind: `{_kind_of(e)}`")
                addr = _crystal_addr_of(e)
                cidx = _crystal_index_of(e)
                if addr:
                    meta.append(f"- crystal_address: `{addr}`")
                if cidx is not None:
                    meta.append(f"- crystal_index: `{cidx}`")
                if e.tags:
                    meta.append(f"- tags: {', '.join(sorted(set(e.tags)))}")
                if e.links:
                    meta.append("- links:")
                    for lk, hv in sorted(e.links.items()):
                        if not hv:
                            continue
                        meta.append(f"  - {lk}: `{hv}`")

                meta.append(f"- created_at: `{e.created_at}`")
                if e.updated_at:
                    meta.append(f"- updated_at: `{e.updated_at}`")

                lines.extend(meta)

                # Graph edges (outgoing)
                if config.include_graph_edges:
                    try:
                        out_edges = self.store.graph.edges(src=h, limit=50)
                        if out_edges:
                            lines.append("- graph_edges:")
                            for edge in out_edges:
                                note = f" ({edge.note})" if edge.note else ""
                                lines.append(f"  - {edge.relation} → `{edge.dst}`{note}")
                    except Exception:
                        pass

                lines.append("")

                # Body
                body = _sanitize_entry_body(e)
                lines.append(body.rstrip())
                lines.append("")

        # Tag index
        if config.include_tag_index:
            lines.append("## Tag index")
            lines.append("")
            for tag in sorted(tag_index.keys(), key=lambda s: s.lower()):
                lines.append(f"### {tag}")
                for title, h in tag_index[tag]:
                    lines.append(f"- [{title}](#{h})")
                lines.append("")

        return "\n".join(lines).rstrip() + "\n"

    # ------------------------------------------------------------------
    # Export helpers
    # ------------------------------------------------------------------

    def export_markdown(self, output_path: str | os.PathLike[str], config: AtlasBookConfig) -> str:
        md = self.build_markdown(config)
        p = Path(output_path).expanduser().resolve()
        p.parent.mkdir(parents=True, exist_ok=True)
        p.write_text(md, encoding="utf-8")
        return str(p)

    def export_pdf(self, output_path: str | os.PathLike[str], config: AtlasBookConfig) -> str:
        """Export a simple PDF.

        Notes
        -----
        This is intentionally a *minimal* renderer. It supports:
        - headings (#, ##, ###)
        - bullet lists
        - fenced code blocks
        - paragraphs
        """

        md = self.build_markdown(config)

        # Lazy import: reportlab is optional in some environments.
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Preformatted, ListFlowable, ListItem
        from xml.sax.saxutils import escape

        styles = getSampleStyleSheet()
        h1 = styles["Heading1"]
        h2 = styles["Heading2"]
        h3 = styles["Heading3"]
        body = styles["BodyText"]

        doc = SimpleDocTemplate(str(Path(output_path).expanduser().resolve()), pagesize=letter)
        story: List[Any] = []

        lines = md.splitlines()
        i = 0
        in_code = False
        code_lines: List[str] = []

        def flush_paragraph(buf: List[str]):
            if not buf:
                return
            text = " ".join([b.strip() for b in buf if b.strip()])
            if text:
                story.append(Paragraph(escape(text), body))
                story.append(Spacer(1, 8))
            buf.clear()

        paragraph_buf: List[str] = []
        bullet_buf: List[str] = []

        def flush_bullets():
            nonlocal bullet_buf
            if not bullet_buf:
                return
            items = []
            for b in bullet_buf:
                b = b.lstrip("-* ").strip()
                if b:
                    items.append(ListItem(Paragraph(escape(b), body)))
            if items:
                story.append(ListFlowable(items, bulletType='bullet'))
                story.append(Spacer(1, 8))
            bullet_buf = []

        while i < len(lines):
            line = lines[i]

            # Code fences
            if line.strip().startswith("```"):
                if in_code:
                    # close
                    in_code = False
                    flush_paragraph(paragraph_buf)
                    flush_bullets()
                    story.append(Preformatted("\n".join(code_lines), styles["Code"]))
                    story.append(Spacer(1, 8))
                    code_lines = []
                else:
                    in_code = True
                    code_lines = []
                i += 1
                continue

            if in_code:
                code_lines.append(line)
                i += 1
                continue

            # Headings
            if line.startswith("# "):
                flush_paragraph(paragraph_buf)
                flush_bullets()
                story.append(Paragraph(escape(line[2:].strip()), h1))
                story.append(Spacer(1, 10))
                i += 1
                continue
            if line.startswith("## "):
                flush_paragraph(paragraph_buf)
                flush_bullets()
                story.append(Paragraph(escape(line[3:].strip()), h2))
                story.append(Spacer(1, 10))
                i += 1
                continue
            if line.startswith("### "):
                flush_paragraph(paragraph_buf)
                flush_bullets()
                story.append(Paragraph(escape(line[4:].strip()), h3))
                story.append(Spacer(1, 8))
                i += 1
                continue

            # HTML anchors from markdown build; skip
            if line.strip().startswith("<a "):
                i += 1
                continue

            # Bullets
            if line.strip().startswith("-") or line.strip().startswith("*"):
                flush_paragraph(paragraph_buf)
                bullet_buf.append(line.strip())
                i += 1
                # Continue collecting consecutive bullet lines.
                while i < len(lines) and (lines[i].strip().startswith("-") or lines[i].strip().startswith("*")):
                    bullet_buf.append(lines[i].strip())
                    i += 1
                flush_bullets()
                continue

            # Blank line
            if not line.strip():
                flush_paragraph(paragraph_buf)
                flush_bullets()
                i += 1
                continue

            paragraph_buf.append(line)
            i += 1

        flush_paragraph(paragraph_buf)
        flush_bullets()

        doc.build(story)
        return str(Path(output_path).expanduser().resolve())

    def export_docx(self, output_path: str | os.PathLike[str], config: AtlasBookConfig) -> str:
        """Export a basic DOCX from the Markdown build.

        This is a minimal renderer intended for portability.
        It supports headings, bullet lists, code blocks, and paragraphs.
        """
        md = self.build_markdown(config)

        from docx import Document  # type: ignore
        from docx.shared import Pt  # type: ignore

        doc = Document()

        lines = md.splitlines()
        in_code = False
        code_buf: List[str] = []
        para_buf: List[str] = []
        bullet_buf: List[str] = []

        def flush_para():
            nonlocal para_buf
            text = " ".join([p.strip() for p in para_buf if p.strip()]).strip()
            if text:
                doc.add_paragraph(text)
            para_buf = []

        def flush_bullets():
            nonlocal bullet_buf
            for b in bullet_buf:
                b = b.lstrip("-* ").strip()
                if b:
                    doc.add_paragraph(b, style="List Bullet")
            bullet_buf = []

        def flush_code():
            nonlocal code_buf
            if not code_buf:
                return
            # Use a paragraph with monospace font.
            p = doc.add_paragraph()
            run = p.add_run("\n".join(code_buf))
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            code_buf = []

        for line in lines:
            # Ignore HTML anchors used for internal linking in Markdown.
            if line.strip().startswith("<a ") and "id=" in line:
                continue

            # Code fence
            if line.strip().startswith("```"):
                flush_para()
                flush_bullets()
                if in_code:
                    flush_code()
                    in_code = False
                else:
                    in_code = True
                    code_buf = []
                continue

            if in_code:
                code_buf.append(line)
                continue

            # Headings
            if line.startswith("# "):
                flush_para()
                flush_bullets()
                doc.add_heading(line[2:].strip(), level=1)
                continue
            if line.startswith("## "):
                flush_para()
                flush_bullets()
                doc.add_heading(line[3:].strip(), level=2)
                continue
            if line.startswith("### "):
                flush_para()
                flush_bullets()
                doc.add_heading(line[4:].strip(), level=3)
                continue
            if line.startswith("#### "):
                flush_para()
                flush_bullets()
                doc.add_heading(line[5:].strip(), level=4)
                continue

            # Bullet
            if line.strip().startswith(("-", "*")) and (line.strip().startswith("- ") or line.strip().startswith("* ")):
                flush_para()
                bullet_buf.append(line)
                continue

            # Blank line
            if not line.strip():
                flush_para()
                flush_bullets()
                continue

            # Normal paragraph line
            para_buf.append(line)

        flush_para()
        flush_bullets()
        if in_code:
            flush_code()

        p = Path(output_path).expanduser().resolve()
        p.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(p))
        return str(p)

    def export_tex(self, output_path: str | os.PathLike[str], config: AtlasBookConfig) -> str:
        """Export a LaTeX (.tex) document from the Markdown build."""
        tex = self.build_latex(config)
        p = Path(output_path).expanduser().resolve()
        p.parent.mkdir(parents=True, exist_ok=True)
        p.write_text(tex, encoding="utf-8")
        return str(p)

    def build_latex(self, config: AtlasBookConfig) -> str:
        """Build a LaTeX document from the Markdown build.

        This is a pragmatic converter designed for internal export.
        """
        md = self.build_markdown(config)

        def esc(s: str) -> str:
            # Escape LaTeX special chars for normal text (not verbatim).
            repl = {
                "\\": r"\textbackslash{}",
                "&": r"\&",
                "%": r"\%",
                "$": r"\$",
                "#": r"\#",
                "_": r"\_",
                "{": r"\{",
                "}": r"\}",
                "~": r"\textasciitilde{}",
                "^": r"\textasciicircum{}",
            }
            out = s
            # Backslash first to avoid double-escaping.
            out = out.replace("\\", repl["\\"])
            for k, v in repl.items():
                if k == "\\":
                    continue
                out = out.replace(k, v)
            return out

        lines = md.splitlines()
        out_lines: List[str] = []

        # Preamble
        out_lines.append(r"\documentclass[11pt]{article}")
        out_lines.append(r"\usepackage[utf8]{inputenc}")
        out_lines.append(r"\usepackage[T1]{fontenc}")
        out_lines.append(r"\usepackage{hyperref}")
        out_lines.append(r"\usepackage{geometry}")
        out_lines.append(r"\usepackage{enumitem}")
        out_lines.append(r"\geometry{margin=1in}")
        out_lines.append(r"\setlist[itemize]{leftmargin=*}")
        out_lines.append(r"\begin{document}")
        out_lines.append("")

        in_code = False
        in_itemize = False
        para_buf: List[str] = []

        def flush_para():
            nonlocal para_buf
            text = " ".join([p.strip() for p in para_buf if p.strip()]).strip()
            if text:
                out_lines.append(esc(text))
                out_lines.append("")
            para_buf = []

        def close_itemize():
            nonlocal in_itemize
            if in_itemize:
                out_lines.append(r"\end{itemize}")
                out_lines.append("")
                in_itemize = False

        for line in lines:
            # Ignore HTML anchors
            if line.strip().startswith("<a ") and "id=" in line:
                continue

            if line.strip().startswith("```"):
                flush_para()
                close_itemize()
                if in_code:
                    out_lines.append(r"\end{verbatim}")
                    out_lines.append("")
                    in_code = False
                else:
                    out_lines.append(r"\begin{verbatim}")
                    in_code = True
                continue

            if in_code:
                out_lines.append(line)
                continue

            # Headings
            if line.startswith("# "):
                flush_para()
                close_itemize()
                out_lines.append(r"\section{" + esc(line[2:].strip()) + "}")
                out_lines.append("")
                continue
            if line.startswith("## "):
                flush_para()
                close_itemize()
                out_lines.append(r"\subsection{" + esc(line[3:].strip()) + "}")
                out_lines.append("")
                continue
            if line.startswith("### "):
                flush_para()
                close_itemize()
                out_lines.append(r"\subsubsection{" + esc(line[4:].strip()) + "}")
                out_lines.append("")
                continue
            if line.startswith("#### "):
                flush_para()
                close_itemize()
                out_lines.append(r"\paragraph{" + esc(line[5:].strip()) + "}")
                out_lines.append("")
                continue

            # Bullet lists
            if line.strip().startswith("- ") or line.strip().startswith("* "):
                flush_para()
                if not in_itemize:
                    out_lines.append(r"\begin{itemize}")
                    in_itemize = True
                item = line.strip()[2:].strip()
                out_lines.append(r"\item " + esc(item))
                continue

            # Blank line
            if not line.strip():
                flush_para()
                close_itemize()
                continue

            para_buf.append(line)

        flush_para()
        close_itemize()
        if in_code:
            out_lines.append(r"\end{verbatim}")
            out_lines.append("")

        out_lines.append(r"\end{document}")
        out_lines.append("")
        return "\n".join(out_lines)
